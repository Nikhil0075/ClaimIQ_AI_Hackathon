"""
ClaimIQ — Attachment Handler
==============================
Three responsibilities:
  1. EXTRACT  — pull attachment bytes out of a raw MIME email message
  2. ANALYZE  — use OpenAI to analyze attachment metadata/text for claim facts
  3. UPLOAD   — save all files to Google Drive under ClaimIQ Claims/{claim_id}/

Google Drive auth: OAuth client JSON
  Set GOOGLE_OAUTH_CREDENTIALS to an OAuth client JSON file.
  First run opens a browser for consent and caches the token under ~/.claimiq/.

Supported file types:
  PDFs/images   → metadata captured, uploaded to Drive, and summarized for agents
  Text/CSV      → read as plain text
  Other         → uploaded to Drive but not analyzed

Usage:
  attachments = extract(raw_email_bytes)
  summary     = analyze(attachments)          # before agents
  folder_url  = upload_to_drive(claim_id, attachments)   # after pipeline
"""

import email as _email_lib
import io
import json
import logging
import mimetypes
import os
import re
import shutil
from typing import Optional

try:
    from dotenv import load_dotenv
except ImportError:
    def load_dotenv(*args, **kwargs):
        return False
from claimiq.shared.openai_client import generate_json, generate_json_messages, image_content_part
from claimiq.shared.config import settings

load_dotenv(dotenv_path=os.path.join(os.path.dirname(__file__), "..", ".env"))

log = logging.getLogger("claimiq.attachments")

DRIVE_ROOT_FOLDER_NAME = os.getenv("DRIVE_ROOT_FOLDER_NAME", "ClaimIQ Claims")
ATTACHMENT_TEXT_MODEL = os.getenv("CLAIMIQ_ATTACHMENT_TEXT_MODEL", settings.lightweight_model)
ATTACHMENT_VISION_MODEL = os.getenv("CLAIMIQ_ATTACHMENT_VISION_MODEL", settings.vision_model)
ATTACHMENT_SYNTHESIS_MODEL = os.getenv("CLAIMIQ_ATTACHMENT_SYNTHESIS_MODEL", settings.lightweight_model)

# Mime types that can be summarized for agent context
ANALYZABLE_IMAGE_TYPES = {"image/jpeg", "image/png", "image/webp", "image/heic", "image/gif"}
ANALYZABLE_PDF_TYPE    = "application/pdf"
ANALYZABLE_TEXT_TYPES  = {"text/plain", "text/csv"}
ANALYZABLE_DOCX_TYPES  = {"application/vnd.openxmlformats-officedocument.wordprocessingml.document"}

# ── Data type ─────────────────────────────────────────────────────────────────

class Attachment:
    """Holds a single email attachment."""
    def __init__(self, filename: str, data: bytes, mime_type: str):
        self.filename  = filename
        self.data      = data
        self.mime_type = mime_type or _guess_mime(filename)
        self.size      = len(data)

    def is_image(self) -> bool:
        return self.mime_type in ANALYZABLE_IMAGE_TYPES

    def is_pdf(self) -> bool:
        return self.mime_type == ANALYZABLE_PDF_TYPE

    def is_text(self) -> bool:
        return self.mime_type in ANALYZABLE_TEXT_TYPES

    def is_docx(self) -> bool:
        return self.mime_type in ANALYZABLE_DOCX_TYPES or self.filename.lower().endswith(".docx")

    def is_analyzable(self) -> bool:
        return self.is_image() or self.is_pdf() or self.is_text() or self.is_docx()

    def __repr__(self):
        return f"<Attachment {self.filename} {self.size//1024}KB {self.mime_type}>"


def _guess_mime(filename: str) -> str:
    t, _ = mimetypes.guess_type(filename)
    return t or "application/octet-stream"


# ── 1. EXTRACT ────────────────────────────────────────────────────────────────

def extract(raw_email_bytes: bytes) -> list[Attachment]:
    """
    Parse raw RFC-822 email bytes and extract all non-body attachments.
    Returns list of Attachment objects (may be empty).
    """
    msg = _email_lib.message_from_bytes(raw_email_bytes)
    found = []

    def _walk(part):
        filename = part.get_filename()
        if not filename:
            return
        ctype = part.get_content_type()
        payload = part.get_payload(decode=True)
        if payload is None:
            return

        safe_name = _safe_filename(filename)
        found.append(Attachment(safe_name, payload, ctype))
        log.info("[Attachments] Extracted: %s (%dKB)", safe_name, len(payload) // 1024)

    for part in msg.walk():
        _walk(part)

    log.info("[Attachments] Total extracted: %d file(s)", len(found))
    return found


def _safe_filename(name: str) -> str:
    """Sanitize attachment filename."""
    # Decode MIME-encoded filenames
    parts = _email_lib.header.decode_header(name)
    decoded = ""
    for chunk, enc in parts:
        if isinstance(chunk, bytes):
            decoded += chunk.decode(enc or "utf-8", errors="replace")
        else:
            decoded += chunk
    return re.sub(r"[^\w.\-]", "_", decoded.strip())


# ── 2. ANALYZE ────────────────────────────────────────────────────────────────

ANALYZE_SYSTEM = """
You are a document analyst for ClaimIQ, an insurance claims AI system.
You will be shown one or more claim-related documents (invoices, police reports,
medical records, photos, etc.). Extract all facts relevant to processing the claim.
Be precise and factual. Flag any inconsistencies or suspicious details.
"""

ANALYZE_USER = """
Analyze the following insurance claim document: {filename}
Document type context: {doc_type}

Extract all relevant facts and return ONLY valid JSON:
{{
  "filename":       "{filename}",
  "document_type":  "<invoice|police_report|medical_record|photo|fir|estimate|id_proof|other>",
  "summary":        "<2-3 sentence summary of what this document shows>",
  "key_facts": {{
    "dates":           ["<any dates found>"],
    "amounts":         ["<any monetary amounts found>"],
    "names":           ["<people, companies, vendors mentioned>"],
    "reference_numbers":["<FIR numbers, invoice numbers, policy refs, etc.>"],
    "vehicle_details": ["<vehicle make, model, reg numbers if present>"],
    "location_details":["<addresses, locations mentioned>"]
  }},
  "risk_signals":   ["<anything suspicious, inconsistent, or worth flagging>"],
  "supports_claim": <true|false|null>,
  "confidence":     <0.0-1.0>
}}
"""

COMBINED_SUMMARY_PROMPT = """
You are the document coordinator for ClaimIQ. You have analyzed {n} documents.
Synthesize all per-document findings into a single structured summary for the claim agents.

INDIVIDUAL DOCUMENT ANALYSES:
{analyses_json}

Return ONLY valid JSON:
{{
  "total_documents": {n},
  "documents_analyzed": ["<filename list>"],
  "aggregate_summary": "<3-4 sentences covering what all documents together show>",
  "all_dates_found":   ["<deduplicated list of all dates>"],
  "all_amounts_found": ["<deduplicated list of all monetary amounts>"],
  "all_vendors":       ["<deduplicated vendor/company names>"],
  "all_references":    ["<all FIR, invoice, policy reference numbers>"],
  "risk_signals":      ["<all risk signals from all documents>"],
  "documents_supporting_claim": ["<filenames that support the claim>"],
  "documents_contradicting_claim": ["<filenames with inconsistencies>"],
  "missing_documents": ["<documents mentioned in email but not provided>"],
  "analyst_notes":     "<anything the adjuster should know about the document set>"
}}
"""


def analyze(attachments: list[Attachment]) -> dict:
    """
    Use OpenAI to analyze attachment context, then synthesize into one summary.
    Returns a combined document summary dict to be passed to agents.
    Returns empty summary dict if no analyzable attachments.
    """
    if not attachments:
        return _empty_summary()

    analyzable = [a for a in attachments if a.is_analyzable()]
    if not analyzable:
        log.info("[Attachments] No analyzable files (only: %s)", [a.filename for a in attachments])
        return _empty_summary(filenames=[a.filename for a in attachments])

    log.info("[Attachments] Analyzing %d/%d file(s) with OpenAI", len(analyzable), len(attachments))

    per_doc = []
    for att in analyzable:
        result = _analyze_one(att)
        if result:
            per_doc.append(result)

    if not per_doc:
        return _empty_summary(filenames=[a.filename for a in attachments])

    if len(per_doc) == 1:
        # Single doc — wrap it directly
        d = per_doc[0]
        return {
            "total_documents":       len(attachments),
            "documents_analyzed":    [d.get("filename")],
            "aggregate_summary":     d.get("summary", ""),
            "all_dates_found":       d.get("key_facts", {}).get("dates", []),
            "all_amounts_found":     d.get("key_facts", {}).get("amounts", []),
            "all_vendors":           d.get("key_facts", {}).get("names", []),
            "all_references":        d.get("key_facts", {}).get("reference_numbers", []),
            "risk_signals":          d.get("risk_signals", []),
            "documents_supporting_claim":      [d["filename"]] if d.get("supports_claim") else [],
            "documents_contradicting_claim":   [],
            "missing_documents":     [],
            "analyst_notes":         "",
            "per_document":          per_doc,
            "modalities":            _modalities(attachments),
        }

    # Multiple docs — synthesize
    return _synthesize(per_doc, len(attachments))


def _analyze_one(att: Attachment) -> Optional[dict]:
    """Analyze a single attachment with OpenAI. Returns dict or fallback metadata."""
    try:
        doc_type = (
            "image/photo" if att.is_image()
            else "PDF document" if att.is_pdf()
            else "Word document" if att.is_docx()
            else "text document"
        )

        prompt_text = ANALYZE_USER.format(
            filename=att.filename,
            doc_type=doc_type,
        )

        if att.is_text():
            text_content = att.data.decode("utf-8", errors="replace")[:4000]
            prompt = f"DOCUMENT: {att.filename}\n\n{text_content}\n\n{prompt_text}"
            result = generate_json(prompt, temperature=0.05, max_tokens=4096, model=ATTACHMENT_TEXT_MODEL)
        elif att.is_image():
            result = _analyze_image(att, prompt_text)
        elif att.is_pdf():
            text_content = _extract_pdf_text(att)
            if text_content:
                prompt = f"PDF DOCUMENT: {att.filename}\n\n{text_content[:6000]}\n\n{prompt_text}"
                result = generate_json(prompt, temperature=0.05, max_tokens=4096, model=ATTACHMENT_TEXT_MODEL)
            else:
                result = _analyze_pdf_image(att, prompt_text) or _fallback_visual_metadata(att, "scanned_pdf")
        elif att.is_docx():
            text_content = _extract_docx_text(att)
            if text_content:
                prompt = f"WORD DOCUMENT: {att.filename}\n\n{text_content[:6000]}\n\n{prompt_text}"
                result = generate_json(prompt, temperature=0.05, max_tokens=4096, model=ATTACHMENT_TEXT_MODEL)
            else:
                result = _fallback_visual_metadata(att, "docx")
        else:
            result = _fallback_visual_metadata(att, doc_type)
        log.info("[Attachments] Analyzed: %s | type=%s | supports=%s",
                 att.filename, result.get("document_type"), result.get("supports_claim"))
        return result

    except Exception as exc:
        log.error("[Attachments] Analysis failed for %s: %s", att.filename, exc)
        modality = "text" if att.is_text() else "pdf" if att.is_pdf() else "image" if att.is_image() else "docx" if att.is_docx() else "other"
        return {
            "filename":      att.filename,
            "document_type": modality,
            "summary":       f"Analysis failed: {exc}",
            "key_facts":     {},
            "risk_signals":  [],
            "supports_claim": None,
            "confidence":    0.0,
            "modality":      modality,
            "mime_type":     att.mime_type,
            "error":         str(exc),
        }


def _analyze_image(att: Attachment, prompt_text: str) -> dict:
    messages = [
        {"role": "system", "content": ANALYZE_SYSTEM},
        {
            "role": "user",
            "content": [
                {"type": "text", "text": prompt_text},
                image_content_part(att.data, att.mime_type),
            ],
        },
    ]
    result = generate_json_messages(messages, temperature=0.05, max_tokens=4096, model=ATTACHMENT_VISION_MODEL)
    result.setdefault("modality", "image")
    result.setdefault("mime_type", att.mime_type)
    return result


def _analyze_pdf_image(att: Attachment, prompt_text: str) -> Optional[dict]:
    extracted = _extract_first_pdf_image(att.data)
    if not extracted:
        return None
    image_data, image_mime_type = extracted
    messages = [
        {"role": "system", "content": ANALYZE_SYSTEM},
        {
            "role": "user",
            "content": [
                {
                    "type": "text",
                    "text": (
                        f"{prompt_text}\n\n"
                        "This page was extracted from a scanned or image-based PDF. "
                        "Read visible text and visual evidence directly from the image."
                    ),
                },
                image_content_part(image_data, image_mime_type),
            ],
        },
    ]
    result = generate_json_messages(messages, temperature=0.05, max_tokens=4096, model=ATTACHMENT_VISION_MODEL)
    result.setdefault("modality", "pdf_image")
    result.setdefault("mime_type", att.mime_type)
    result.setdefault("rendered_from_pdf", True)
    return result


def _extract_pdf_text(att: Attachment) -> str:
    """Best-effort PDF text extraction without making pypdf a hard dependency."""
    try:
        from pypdf import PdfReader

        reader = PdfReader(io.BytesIO(att.data))
        return "\n".join((page.extract_text() or "") for page in reader.pages[:5]).strip()
    except Exception as exc:
        log.info("[Attachments] PDF text extraction unavailable for %s: %s", att.filename, exc)
        return ""


def _extract_first_pdf_image(data: bytes) -> Optional[tuple[bytes, str]]:
    try:
        from pypdf import PdfReader

        reader = PdfReader(io.BytesIO(data))
        for page in reader.pages[:3]:
            images = list(getattr(page, "images", []) or [])
            if not images:
                continue
            image = max(images, key=lambda item: len(getattr(item, "data", b"") or b""))
            image_data = getattr(image, "data", b"") or b""
            if not image_data:
                continue
            name = str(getattr(image, "name", "") or "").lower()
            if name.endswith((".jpg", ".jpeg")):
                return image_data, "image/jpeg"
            if name.endswith(".png"):
                return image_data, "image/png"
            return image_data, "image/jpeg"
    except Exception as exc:
        log.info("[Attachments] PDF image extraction unavailable: %s", exc)
        return None
    return None


def _extract_docx_text(att: Attachment) -> str:
    """Best-effort DOCX text extraction without making python-docx a hard dependency."""
    try:
        from docx import Document

        doc = Document(io.BytesIO(att.data))
        paragraphs = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]
        table_cells = []
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        table_cells.append(cell.text)
        return "\n".join(paragraphs + table_cells).strip()
    except Exception as exc:
        log.info("[Attachments] DOCX text extraction unavailable for %s: %s", att.filename, exc)
        return ""


def _fallback_visual_metadata(att: Attachment, doc_type: str) -> dict:
    modality = "pdf" if att.is_pdf() else "image" if att.is_image() else "docx" if att.is_docx() else "other"
    return {
        "filename": att.filename,
        "document_type": doc_type,
        "summary": f"{doc_type} received. Text/visual extraction was not available; preserve as evidence and review manually.",
        "key_facts": {
            "dates": [],
            "amounts": [],
            "names": [],
            "reference_numbers": [],
            "vehicle_details": [],
            "location_details": [],
        },
        "risk_signals": ["manual_review_needed_for_unreadable_document"] if doc_type == "scanned_pdf" else [],
        "supports_claim": None,
        "confidence": 0.25,
        "modality": modality,
        "mime_type": att.mime_type,
        "size_bytes": att.size,
    }


def _synthesize(per_doc: list[dict], total_count: int) -> dict:
    """Combine multiple document analyses into one summary via OpenAI."""
    try:
        prompt = COMBINED_SUMMARY_PROMPT.format(
            n=len(per_doc),
            analyses_json=json.dumps(per_doc, indent=2, default=str),
        )

        result = generate_json(prompt, temperature=0.05, max_tokens=4096, model=ATTACHMENT_SYNTHESIS_MODEL)
        result["per_document"] = per_doc
        result["total_documents"] = total_count
        result["modalities"] = _per_doc_modalities(per_doc)
        log.info("[Attachments] Synthesis complete | risk_signals=%d", len(result.get("risk_signals", [])))
        return result

    except Exception as exc:
        log.error("[Attachments] Synthesis failed: %s", exc)
        # Fall back to manual merge
        return _manual_merge(per_doc, total_count)


def _manual_merge(per_doc: list[dict], total_count: int) -> dict:
    """Fallback merge when OpenAI synthesis fails."""
    dates, amounts, vendors, refs, risks, supporting = [], [], [], [], [], []
    for d in per_doc:
        kf = d.get("key_facts", {})
        dates   += kf.get("dates", [])
        amounts += kf.get("amounts", [])
        vendors += kf.get("names", [])
        refs    += kf.get("reference_numbers", [])
        risks   += d.get("risk_signals", [])
        if d.get("supports_claim"):
            supporting.append(d.get("filename", ""))

    return {
        "total_documents":       total_count,
        "documents_analyzed":    [d.get("filename") for d in per_doc],
        "aggregate_summary":     " | ".join(d.get("summary", "") for d in per_doc),
        "all_dates_found":       list(dict.fromkeys(dates)),
        "all_amounts_found":     list(dict.fromkeys(amounts)),
        "all_vendors":           list(dict.fromkeys(vendors)),
        "all_references":        list(dict.fromkeys(refs)),
        "risk_signals":          list(dict.fromkeys(risks)),
        "documents_supporting_claim":    supporting,
        "documents_contradicting_claim": [],
        "missing_documents":     [],
        "analyst_notes":         "Synthesis fallback — review per_document for details.",
        "per_document":          per_doc,
        "modalities":            _per_doc_modalities(per_doc),
    }


def _empty_summary(filenames: list[str] = None) -> dict:
    return {
        "total_documents":       len(filenames or []),
        "documents_analyzed":    [],
        "aggregate_summary":     "No analyzable attachments provided.",
        "all_dates_found":       [],
        "all_amounts_found":     [],
        "all_vendors":           [],
        "all_references":        [],
        "risk_signals":          [],
        "documents_supporting_claim":    [],
        "documents_contradicting_claim": [],
        "missing_documents":     [],
        "analyst_notes":         "",
        "per_document":          [],
        "modalities":            [],
    }


def _modalities(attachments: list[Attachment]) -> list[str]:
    values = []
    for att in attachments:
        if att.is_image():
            values.append("image")
        elif att.is_pdf():
            values.append("pdf")
        elif att.is_text():
            values.append("text")
        elif att.is_docx():
            values.append("docx")
        else:
            values.append(att.mime_type)
    return sorted(set(values))


def _per_doc_modalities(per_doc: list[dict]) -> list[str]:
    return sorted(set(
        str(d.get("modality") or d.get("document_type"))
        for d in per_doc
        if d.get("modality") or d.get("document_type")
    ))


# ── 3. UPLOAD TO GOOGLE DRIVE ─────────────────────────────────────────────────

DRIVE_SCOPES = ["https://www.googleapis.com/auth/drive.file"]

# Token stored locally so OAuth only runs once
DRIVE_TOKEN_PATH = os.path.join(
    os.path.expanduser("~"), ".claimiq", "drive_token.json"
)
DRIVE_TOKEN_PATH = os.getenv("GOOGLE_DRIVE_TOKEN_PATH", DRIVE_TOKEN_PATH)
CREDENTIALS_JSON = os.getenv(
    "GOOGLE_OAUTH_CREDENTIALS",
    os.path.join(os.path.dirname(__file__), "..", "credentials.json"),
)

_drive_service = None


def _is_invalid_grant(exc: Exception) -> bool:
    text = str(exc).lower()
    return "invalid_grant" in text or "token has been expired or revoked" in text


def _clear_cached_drive_token(reason: str) -> None:
    """Move a rejected OAuth token aside so Drive can re-authorize."""
    if not os.path.exists(DRIVE_TOKEN_PATH):
        return
    backup_path = f"{DRIVE_TOKEN_PATH}.invalid"
    try:
        if os.path.exists(backup_path):
            os.remove(backup_path)
        shutil.move(DRIVE_TOKEN_PATH, backup_path)
        log.warning("[Drive] Cached OAuth token rejected (%s); moved it to %s", reason, backup_path)
    except Exception as exc:
        log.warning("[Drive] Could not clear rejected OAuth token %s: %s", DRIVE_TOKEN_PATH, exc)


def _truthy_env(name: str) -> bool | None:
    value = os.getenv(name)
    if value is None:
        return None
    return value.strip().lower() in {"1", "true", "yes", "on"}


def _browser_oauth_allowed() -> bool:
    """Return whether this runtime should attempt an interactive browser OAuth flow."""
    override = _truthy_env("CLAIMIQ_ALLOW_BROWSER_OAUTH")
    if override is not None:
        return override
    if os.getenv("CLAIMIQ_DRIVE_TOKEN_SOURCE") or os.getenv("GOOGLE_DRIVE_TOKEN_PATH"):
        return False
    if os.name != "nt" and not (
        os.getenv("DISPLAY") or os.getenv("WAYLAND_DISPLAY") or os.getenv("BROWSER")
    ):
        return False
    return True


def _browser_oauth_unavailable_message() -> str:
    return (
        "Drive OAuth token is missing, expired, or revoked, and this runtime cannot "
        "open a browser for a new consent flow. Regenerate the Drive token locally "
        "with the same GOOGLE_OAUTH_CREDENTIALS client, then update the Streamlit "
        "secret GOOGLE_DRIVE_TOKEN_JSON or GOOGLE_DRIVE_REFRESH_TOKEN. If the Google "
        "OAuth consent app is in Testing mode, publish it to Production or refresh "
        "tokens may expire/revoke unexpectedly."
    )


def _get_drive_service():
    """
    Build and cache the Drive API service using OAuth user credentials.
    OAuth uploads files into the user's personal Drive.

    First run: opens a browser for one-time authorization.
    Subsequent runs: uses the cached token at ~/.claimiq/drive_token.json.
    """
    global _drive_service
    if _drive_service is not None:
        return _drive_service

    try:
        from googleapiclient.discovery import build
        from google.oauth2.credentials import Credentials
        from google_auth_oauthlib.flow import InstalledAppFlow
        from google.auth.transport.requests import Request

        creds = None

        # Load cached token
        os.makedirs(os.path.dirname(DRIVE_TOKEN_PATH), exist_ok=True)
        if os.path.exists(DRIVE_TOKEN_PATH):
            creds = Credentials.from_authorized_user_file(DRIVE_TOKEN_PATH, DRIVE_SCOPES)

        # Refresh or run browser flow
        if not creds or not creds.valid:
            if creds and creds.expired and creds.refresh_token:
                log.info("[Drive] Refreshing OAuth token...")
                try:
                    creds.refresh(Request())
                    log.info("[Drive] OAuth token refreshed successfully")
                except Exception as exc:
                    if not _is_invalid_grant(exc):
                        raise
                    _clear_cached_drive_token("invalid_grant during refresh")
                    creds = None

            if not creds or not creds.valid:
                if not _browser_oauth_allowed():
                    raise RuntimeError(_browser_oauth_unavailable_message())
                if not os.path.exists(CREDENTIALS_JSON):
                    raise FileNotFoundError(
                        f"credentials.json not found at {CREDENTIALS_JSON}. "
                        "Download it from Google Cloud Console -> APIs & Services -> Credentials -> OAuth 2.0 Client."
                    )
                log.info("[Drive] Opening browser for OAuth authorization (one-time)...")
                flow = InstalledAppFlow.from_client_secrets_file(CREDENTIALS_JSON, DRIVE_SCOPES)
                creds = flow.run_local_server(port=0, open_browser=True)

            # Save for next run
            with open(DRIVE_TOKEN_PATH, "w") as f:
                f.write(creds.to_json())
            log.info("[Drive] Token cached at %s", DRIVE_TOKEN_PATH)

        _drive_service = build("drive", "v3", credentials=creds, cache_discovery=False)
        log.info("[Drive] Service initialized (OAuth user credentials)")
        return _drive_service

    except Exception as exc:
        log.error("[Drive] Failed to initialize Drive service: %s", exc)
        raise


def _get_or_create_folder(service, name: str, parent_id: Optional[str] = None) -> str:
    """Get an existing folder by name (under parent) or create it. Returns folder ID."""
    query = f"name='{name}' and mimeType='application/vnd.google-apps.folder' and trashed=false"
    if parent_id:
        query += f" and '{parent_id}' in parents"

    results = service.files().list(
        q=query, spaces="drive", fields="files(id, name)", pageSize=1
    ).execute()

    files = results.get("files", [])
    if files:
        return files[0]["id"]

    # Create it
    metadata = {
        "name":     name,
        "mimeType": "application/vnd.google-apps.folder",
    }
    if parent_id:
        metadata["parents"] = [parent_id]

    folder = service.files().create(body=metadata, fields="id").execute()
    log.info("[Drive] Created folder: %s", name)
    return folder["id"]


def upload_file_to_drive(claim_id: str, filename: str, data: bytes, mime_type: str = "application/pdf") -> bool:
    """
    Upload a single file (bytes) to the claim's Google Drive folder.
    Creates the folder structure if it doesn't already exist.
    Returns True on success, False on failure.

    Typical use: upload the AI-generated adjuster guide PDF after the main
    attachment batch has been sent, or when there are no other attachments.
    """
    if not data:
        log.warning("[Drive] upload_file_to_drive called with empty data for %s", filename)
        return False
    try:
        from googleapiclient.http import MediaIoBaseUpload

        service = _get_drive_service()
        root_id    = _get_or_create_folder(service, DRIVE_ROOT_FOLDER_NAME)
        claim_folder_id = _get_or_create_folder(service, claim_id, parent_id=root_id)

        metadata = {"name": filename, "parents": [claim_folder_id]}
        media = MediaIoBaseUpload(io.BytesIO(data), mimetype=mime_type, resumable=False)
        service.files().create(body=metadata, media_body=media, fields="id").execute()

        # Ensure folder is readable (idempotent — silently ignored if already set)
        try:
            service.permissions().create(
                fileId=claim_folder_id,
                body={"role": "reader", "type": "anyone"},
            ).execute()
        except Exception:
            pass

        log.info("[Drive] Uploaded file: %s to claim folder %s (%d bytes)", filename, claim_id, len(data))
        return True

    except Exception as exc:
        log.error("[Drive] upload_file_to_drive failed for %s / %s: %s", claim_id, filename, exc)
        return False


def upload_to_drive(claim_id: str, attachments: list[Attachment]) -> str:
    """
    Upload all attachments to Google Drive under:
      ClaimIQ Claims / {claim_id} /

    Returns the shareable folder URL (viewer link).
    Returns empty string if Drive is not configured or upload fails.
    """
    if not attachments:
        log.info("[Drive] No attachments to upload")
        return ""

    try:
        from googleapiclient.http import MediaIoBaseUpload

        service = _get_drive_service()

        # Get/create root and claim folders
        root_id  = _get_or_create_folder(service, DRIVE_ROOT_FOLDER_NAME)
        claim_id_folder = _get_or_create_folder(service, claim_id, parent_id=root_id)

        uploaded = 0
        for att in attachments:
            try:
                metadata = {
                    "name":    att.filename,
                    "parents": [claim_id_folder],
                }
                media = MediaIoBaseUpload(
                    io.BytesIO(att.data),
                    mimetype=att.mime_type,
                    resumable=len(att.data) > 5_000_000,  # resumable for >5MB
                )
                service.files().create(
                    body=metadata,
                    media_body=media,
                    fields="id",
                ).execute()
                log.info("[Drive] Uploaded: %s (%dKB)", att.filename, att.size // 1024)
                uploaded += 1
            except Exception as exc:
                log.error("[Drive] Failed to upload %s: %s", att.filename, exc)

        # Make folder readable by anyone with the link
        try:
            service.permissions().create(
                fileId=claim_id_folder,
                body={"role": "reader", "type": "anyone"},
            ).execute()
        except Exception:
            pass  # Non-critical — folder still accessible to service account

        folder_url = f"https://drive.google.com/drive/folders/{claim_id_folder}"
        log.info("[Drive] Upload complete | %d/%d files | %s", uploaded, len(attachments), folder_url)
        return folder_url

    except Exception as exc:
        log.error("[Drive] Upload failed: %s", exc)
        return ""
